import os
import streamlit as st
import fitz  # PyMuPDF
from io import BytesIO
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# --- 1. GLOBAL CONFIGURATION & AUTH ---
st.set_page_config(page_title="History 9489 Portal", layout="wide")

# Folder setup
FOLDERS = ["9489_June_qp", "9489_Nov_qp", "9489_June_ms", "9489_Nov_ms"]
for folder in FOLDERS:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Google Drive Config
SERVICE_ACCOUNT_FILE = 'credentials.json'
SCOPES = ['https://www.googleapis.com/auth/drive']
FOLDER_ID = '1hA1IXjC5mJEDLnGU0BSpLy2wPgJ-6R3x'

# Initialize session states
if 'basket' not in st.session_state:
    st.session_state.basket = []
if 'search_results' not in st.session_state:
    st.session_state.search_results = []


# --- 2. CORE FUNCTIONS (Must be at the top) ---

def get_gdrive_service():
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    return build('drive', 'v3', credentials=creds)


def sync_from_drive():
    """Teleports files from Google Drive to local folders"""
    service = get_gdrive_service()
    results = service.files().list(
        q=f"'{FOLDER_ID}' in parents and mimeType='application/pdf' and trashed=false",
        fields="files(id, name)").execute()
    items = results.get('files', [])

    if not items:
        st.sidebar.warning("No PDFs found in Cloud.")
        return

    for item in items:
        file_id = item['id']
        file_name = item['name']

        # Determine target folder based on filename
        target = "9489_June_qp"
        name_low = file_name.lower()
        if "ms" in name_low:
            target = "9489_June_ms" if ("june" in name_low or "_s" in name_low) else "9489_Nov_ms"
        elif "nov" in name_low or "_w" in name_low:
            target = "9489_Nov_qp"

        local_path = os.path.join(target, file_name)
        if not os.path.exists(local_path):
            request = service.files().get_media(fileId=file_id)
            fh = BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            with open(local_path, "wb") as f:
                f.write(fh.getvalue())


def search_pdfs(keyword):
    """Scans local PDFs for keyword text"""
    results = []
    for folder in FOLDERS:
        if os.path.exists(folder):
            for file in os.listdir(folder):
                if file.endswith(".pdf"):
                    file_path = os.path.join(folder, file)
                    try:
                        doc = fitz.open(file_path)
                        found_in_file = False
                        for page in doc:
                            if keyword.lower() in page.get_text().lower():
                                found_in_file = True
                                break
                        doc.close()
                        if found_in_file:
                            results.append({"name": file, "path": file_path})
                    except Exception as e:
                        print(f"Error reading {file}: {e}")
    return results
################################################################

    st.divider()

    # Job A: Manual Upload Link
    st.subheader("📤 Add New Papers")
    st.info("Upload files to Drive first, then click 'Sync' above.")
    folder_url = "https://drive.google.com/drive/folders/1hA1IXjC5mJEDLnGU0BSpLy2wPgJ-6R3x"
    st.link_button("📂 Open Google Drive Warehouse", folder_url)

    st.divider()
    st.markdown("**Syllabus:** History 9489")
    st.markdown("**Location:** Brunei Darussalam")

# --- 4. MAIN INTERFACE ---
st.title("PUSAT TINGKATAN ENAM SENGKURONG")
st.title("📜 9489 A Level History PYP Portal")
st.markdown("---")

tab1, tab2, tab3 = st.tabs(["🔍 Search Topics", "🧺 Handout Basket", "🔐 Lecturer Hub"])

# --- TAB 1: SEARCH ---
with tab1:
    st.subheader("Search Past Papers")
    keyword = st.text_input("Enter a historical topic (e.g., 'Civil War'):", key="search_input")

    if st.button("Run Search"):
        if keyword:
            with st.spinner('Scanning archives...'):
                st.session_state.search_results = search_pdfs(keyword)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.search_results:
        st.success(f"Found {len(st.session_state.search_results)} relevant papers.")
        for item in st.session_state.search_results:
            col1, col2 = st.columns([4, 1])
            col1.write(f"📄 {item['name']}")

            is_in_basket = any(b['name'] == item['name'] for b in st.session_state.basket)
            if col2.button("➕ Add", key=f"add_{item['name']}"):
                if not is_in_basket:
                    st.session_state.basket.append(item)
                    st.toast(f"Added {item['name']}!")
                else:
                    st.toast("Already in basket!")

# --- TAB 2: BASKET ---
with tab2:
    st.subheader("Your Handout Collection")
    if st.session_state.basket:
        for i, paper in enumerate(st.session_state.basket):
            st.write(f"{i + 1}. {paper['name']}")

        st.divider()
        col_clear, col_word, col_pdf = st.columns([1, 1.5, 1.5])

        with col_clear:
            if st.button("🗑️ Clear All"):
                st.session_state.basket = []
                st.rerun()

        with col_word:
            if st.button("📄 Word (Text Only)"):
                doc = Document()
                doc.add_heading('PTES History 9489 Handout', 0)
                for paper in st.session_state.basket:
                    doc.add_heading(f"Source: {paper['name']}", level=1)
                    pdf = fitz.open(paper['path'])
                    for page in pdf:
                        doc.add_paragraph(page.get_text())
                    pdf.close()
                bio_word = BytesIO()
                doc.save(bio_word)
                st.download_button("📥 Download Word", bio_word.getvalue(), "handout.docx")

        with col_pdf:
            if st.button("📑 PDF (Snipped Pages)"):
                if not keyword:
                    st.error("Enter a keyword in Search Tab first!")
                else:
                    merged_pdf = fitz.open()
                    pages_found = 0
                    for paper in st.session_state.basket:
                        pdf_item = fitz.open(paper['path'])
                        for page_index in range(len(pdf_item)):
                            if keyword.lower() in pdf_item[page_index].get_text().lower():
                                merged_pdf.insert_pdf(pdf_item, from_page=page_index, to_page=page_index)
                                pages_found += 1
                        pdf_item.close()

                    if pages_found > 0:
                        bio_pdf = BytesIO()
                        merged_pdf.save(bio_pdf)
                        st.download_button("📥 Download PDF", bio_pdf.getvalue(), "snipped_handout.pdf")
                    else:
                        st.warning("No specific pages matched for snipping.")
    else:
        st.info("Basket is empty.")

# --- TAB 3: LECTURER HUB ---
with tab3:
    st.subheader("📚 Lecturer Resource Manager")
    password = st.text_input("Enter Admin Password", type="password", key="admin_pw")

    if password == "brunei9489":
        c_yr, c_mo, c_tp = st.columns(3)
        sel_year = c_yr.text_input("📅 Year", placeholder="e.g. 2021")
        sel_month = c_mo.selectbox("🌙 Session", ["All", "June (s)", "Nov (w)"])
        sel_type = c_tp.selectbox("📄 Type", ["All", "Question Paper (QP)", "Mark Scheme (MS)"])

        filtered_list = []
        for folder in FOLDERS:
            if sel_type == "Question Paper (QP)" and "_ms" in folder.lower(): continue
            if sel_type == "Mark Scheme (MS)" and "_qp" in folder.lower(): continue
            if sel_month == "June (s)" and "nov" in folder.lower(): continue
            if sel_month == "Nov (w)" and "june" in folder.lower(): continue

            if os.path.exists(folder):
                for file in os.listdir(folder):
                    if not sel_year or sel_year.lower() in file.lower():
                        filtered_list.append({"folder": folder, "name": file})

        for item in filtered_list:
            with st.container():
                col_a, col_b, col_c = st.columns([6, 2, 2])
                file_path = os.path.join(item['folder'], item['name'])
                col_a.write(f"📁 {item['folder']} / {item['name']}")

                with open(file_path, "rb") as f:
                    col_b.download_button("💾 Save", f, file_name=item['name'],
                                          key=f"save_{item['name']}_{item['folder']}")

                if col_c.button("🗑️ Delete", key=f"del_{item['name']}"):
                    os.remove(file_path)
                    st.rerun()
            st.divider()
    elif password != "":
        st.error("Incorrect Password")
###############################################################################
# --- FOOTER ---
st.markdown("---")

# Using a single container with centered alignment
st.markdown(
    """
    <div style="text-align: center; width: 100%;">
        <p style="font-size: 20px; font-weight: bold; margin-bottom: 5px;">
            ✨ PTES 9489 History PYP Resources Portal ✨
        </p>
        <p style="font-size: 16px; font-weight: bold; letter-spacing: 0.5px;">
            <span style="color: #FF0000;">🔴 Academically Excellence</span> | 
            <span style="color: #FFD700;">🟡 Future Readiness</span> | 
            <span style="color: #0070FF;">🔵 Digital & Integrity</span> | 
            <span style="color: #28A745;">🟢 Holistic & Growth</span>
        </p>
        <p style="color: gray; font-size: 14px; margin-top: 10px;">
            Creator: Miss Hajah Nurul Haziqah HN (Computer Sci 18/4/2026)
        </p>
    </div>
    """,
    unsafe_allow_html=True
)
